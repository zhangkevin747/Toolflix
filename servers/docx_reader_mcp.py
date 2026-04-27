"""Minimal MCP server: read text content from .docx files via python-docx."""
from mcp.server.fastmcp import FastMCP
from docx import Document

mcp = FastMCP("docx-reader-mcp")


@mcp.tool()
def read_docx(path: str) -> str:
    """Read text content from a Microsoft Word .docx file.

    Args:
        path: Absolute path to a .docx file.

    Returns:
        Full document text, paragraph-joined with newlines. Tables are
        rendered row-by-row with tab-separated cells.
    """
    doc = Document(path)
    out = []
    for p in doc.paragraphs:
        if p.text.strip():
            out.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            out.append("\t".join(cells))
    return "\n".join(out)


if __name__ == "__main__":
    mcp.run()
